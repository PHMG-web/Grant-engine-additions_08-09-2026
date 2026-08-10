import os
import tempfile
import json
import pytest
from docx import Document
from grant_engine.nofo_parser import NOFOParser, NOFOData
from grant_engine.docx_extractor import DocxExtractor

# ==========================================
# NOFO PARSER TESTS
# ==========================================

def test_nofo_parser_clean_text():
    parser = NOFOParser()
    text = "   Hello    World \t  - This – is a test.   "
    cleaned = parser.clean_text(text)
    # The return should be a string, not a method/function pointer
    assert isinstance(cleaned, str)
    assert cleaned == "Hello World - This - is a test."

def test_nofo_parser_extract_text_invalid_input():
    parser = NOFOParser()
    with pytest.raises(RuntimeError) as exc_info:
        parser.extract_text(None)
    assert "PDF extraction failed" in str(exc_info.value)

    with pytest.raises(RuntimeError) as exc_info:
        parser.extract_text("non_existent_file.pdf")
    assert "PDF extraction failed" in str(exc_info.value)

def test_nofo_parser_parse_empty_input():
    parser = NOFOParser()
    with pytest.raises(ValueError) as exc_info:
        parser.parse("")
    assert "Cannot parse empty or null text" in str(exc_info.value)


# ==========================================
# DOCX EXTRACTOR TESTS
# ==========================================

def test_docx_extractor_load_error():
    with pytest.raises(RuntimeError) as exc_info:
        DocxExtractor("non_existent_file.docx")
    assert "Failed to load DOCX file" in str(exc_info.value)

def test_docx_extractor_workflow():
    # Create a temporary DOCX file
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name

    try:
        # Create a DOCX document using python-docx
        doc = Document()
        doc.add_paragraph("This is a simple paragraph with {Organization_Name} and {{Mission}} placeholders.")
        doc.add_paragraph("Here is a malformed unmatched open brace: {unmatched_open and some other text.")
        doc.add_paragraph("And here is an unmatched closing brace: unmatched_close} with empty braces {} inside.")
        doc.add_paragraph("Empty braces with whitespace {   } and double empty {{}}.")

        # Let's add a table
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell 1 with valid {Capacity} placeholder."
        table.cell(0, 1).text = "Cell 2 with malformed {unmatched_table_cell brace."
        table.cell(1, 0).text = "Cell 3 with empty table cell {} braces."
        table.cell(1, 1).text = "Cell 4 plain text."

        doc.save(tmp_path)

        # Initialize the extractor
        extractor = DocxExtractor(tmp_path)

        # 1. Test full text extraction
        full_text = extractor.extract_full_text()
        assert "Organization_Name" in full_text
        assert "Capacity" in full_text

        # 2. Test valid placeholders extraction
        placeholders = extractor.extract_placeholders()
        # Should extract: "Organization_Name", "Mission", "Capacity"
        assert "Organization_Name" in placeholders
        assert "Mission" in placeholders
        assert "Capacity" in placeholders
        # Make sure no malformed placeholders are returned as valid
        assert "unmatched_open" not in placeholders
        assert "unmatched_close" not in placeholders

        # 3. Test malformed tokens detection
        issues = extractor.detect_malformed_tokens()
        # Let's print issues for debug
        print("Detected issues:")
        for issue in issues:
            print(f" - {issue}")

        # Check for paragraph issues
        assert any("unmatched opening brace" in s for m in issues for s in [m.lower()])
        assert any("unmatched closing brace" in s for m in issues for s in [m.lower()])
        assert any("empty placeholder" in s for m in issues for s in [m.lower()])

        # 4. Test extract_all
        all_data = extractor.extract_all()
        assert "text" in all_data
        assert "placeholders" in all_data
        assert "issues" in all_data
        assert len(all_data["placeholders"]) == 3
        assert len(all_data["issues"]) > 0

    finally:
        # Clean up the temporary file
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


def test_docx_extractor_load_sources():
    # 1. Test bytes / stream loading
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name

    try:
        doc = Document()
        doc.add_paragraph("Testing {Placeholder}")
        doc.save(tmp_path)

        with open(tmp_path, "rb") as f:
            content = f.read()

        import io
        stream = io.BytesIO(content)
        # Load from file-like object (BytesIO)
        extractor = DocxExtractor(stream)
        assert "Placeholder" in extractor.extract_placeholders()

        # 2. Test FastAPI UploadFile-like object loading
        class MockUploadFile:
            def __init__(self, file_obj):
                self.file = file_obj

        stream.seek(0)
        mock_file = MockUploadFile(stream)
        extractor_upload = DocxExtractor(mock_file)
        assert "Placeholder" in extractor_upload.extract_placeholders()

    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


def test_docx_extractor_auto_correct():
    # Test text corrections
    extractor = DocxExtractor.__new__(DocxExtractor) # Create uninitialized instance for unit testing method directly
    
    # 1. Test correct_text_braces logic
    text = "This has {unmatched_open and other matched {correct_field} with {}. Also unmatched close}"
    corrected, fixes = extractor.correct_text_braces(text)
    # {unmatched_open -> {unmatched_open} (1 fix)
    # {} -> removed (1 fix)
    # close} -> {close} (1 fix)
    assert fixes == 3
    assert "{unmatched_open}" in corrected
    assert "{close}" in corrected
    assert "{}" not in corrected

    # 2. Test full document auto-correction
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name

    try:
        doc = Document()
        doc.add_paragraph("Para with {Organization_Name open brace.")
        doc.add_paragraph("Para with Capacity} close brace.")
        doc.add_paragraph("Para with empty {} braces.")
        doc.save(tmp_path)

        ext = DocxExtractor(tmp_path)
        # Verify it has malformed issues before correction
        assert len(ext.detect_malformed_tokens()) > 0

        # Run correction
        fixes_made = ext.auto_correct_braces(tmp_path)
        assert fixes_made == 3

        # Re-load and verify it is clean with NO malformed issues and 100% valid placeholders!
        clean_ext = DocxExtractor(tmp_path)
        assert len(clean_ext.detect_malformed_tokens()) == 0
        placeholders = clean_ext.extract_placeholders()
        assert "Organization_Name" in placeholders
        assert "Capacity" in placeholders

    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


# ==========================================
# GRANT CONTEXT & STATE TESTS
# ==========================================

def test_grant_context_type_guards():
    from grant_engine.grant_context import GrantContext
    context = GrantContext()
    
    # Safely update list conversion
    context.update_section("program_design", {
        "Program_Name": "Test Program",
        "Objectives": "Objective 1, Objective 2, Objective 3" # String gets converted to list
    })
    
    assert context.program_design["Program_Name"] == "Test Program"
    assert context.program_design["Objectives"] == ["Objective 1", "Objective 2", "Objective 3"]

    # Trigger TypeError on incorrect structure types
    with pytest.raises(TypeError):
        context.update_section("program_design", "not-a-dict")


# ==========================================
# SECTION GENERATOR TESTS
# ==========================================

def test_section_generator_formatting():
    from grant_engine.section_generator import SectionGenerator
    gen = SectionGenerator()
    
    # 1. Test single and double braces with spacing
    template = "This is a {Single_Brace} and {{Double_Brace}} and {{  Spaced_Brace  }}."
    variables = {
        "Single_Brace": "A",
        "Double_Brace": "B",
        "Spaced_Brace": "C"
    }
    
    res = gen.generate_section(template, variables)
    assert res["generated_text"] == "This is a A and B and C."

    # 2. Test natural list narrative formatting
    list_template = "Objectives: {Objectives}"
    assert gen.generate_section(list_template, {"Objectives": ["Objective 1"]})["generated_text"] == "Objectives: Objective 1"
    assert gen.generate_section(list_template, {"Objectives": ["Obj 1", "Obj 2"]})["generated_text"] == "Objectives: Obj 1 and Obj 2"
    assert gen.generate_section(list_template, {"Objectives": ["Obj 1", "Obj 2", "Obj 3"]})["generated_text"] == "Objectives: Obj 1, Obj 2, and Obj 3"


# ==========================================
# TEMPLATE LOADER VALIDATION TESTS
# ==========================================

def test_template_loader_guards():
    from grant_engine.template_loader import TemplateLoader
    
    # Directory missing raise
    with pytest.raises(FileNotFoundError):
        TemplateLoader("non_existent_templates_dir")


# ==========================================
# DATA LOADER & MALFORMED JSON TESTS
# ==========================================

def test_data_loader_guards():
    from grant_engine.data_loader import DataLoader
    
    loader = DataLoader("non_existent_data_dir")
    with pytest.raises(FileNotFoundError):
        loader.load_all()


# ==========================================
# EXPORTER DOCX & JSON GUARDS
# ==========================================

def test_exporter_safe_nones():
    from grant_engine.exporter import Exporter
    from grant_engine.grant_context import GrantContext
    
    exp = Exporter()
    context = GrantContext()
    
    # Unpopulated values print safely
    assert exp._safe_value(None) == "[Not Populated]"
    assert exp._safe_value([]) == "[Not Populated]"
    
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
        
    try:
        # Export DOCX safely with missing values without exceptions
        saved_path = exp.export_docx(context, tmp_path)
        assert os.path.exists(saved_path)
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


# ==========================================
# WORKFLOW ENGINE INTEGRATION TESTS
# ==========================================

def test_engine_workflow():
    from grant_engine.engine import GrantAutomationEngine
    
    # Create fake templates directory with manifest and template files
    with tempfile.TemporaryDirectory() as temp_dir:
        manifest_data = {
            "templates": [
                {
                    "id": "organizational_profile",
                    "filename": "org_profile.txt",
                    "order": 1,
                    "required_fields": ["Organization_Name"]
                },
                {
                    "id": "budget_narrative",
                    "filename": "budget.txt",
                    "order": 2,
                    "required_fields": ["Budget_Total"]
                }
            ]
        }
        
        # Write manifest.json
        with open(os.path.join(temp_dir, "manifest.json"), "w", encoding="utf-8") as f:
            json.dump(manifest_data, f)
            
        # Write templates
        with open(os.path.join(temp_dir, "org_profile.txt"), "w", encoding="utf-8") as f:
            f.write("Profile: {Organization_Name} with mission {Mission}.")
            
        with open(os.path.join(temp_dir, "budget.txt"), "w", encoding="utf-8") as f:
            f.write("Budget Narrative is {{Budget_Total}}.")

        # Initialize engine
        engine = GrantAutomationEngine(temp_dir)
        
        # Run workflow and assert validation checks run automatically
        context = engine.run()
        
        assert engine.validation_results is not None
        # Since fields are empty initially, errors will exist
        assert "organizational_profile" in engine.validation_results["errors"]
        assert "budget_narrative" in engine.validation_results["errors"]


# ==========================================
# GRANT SCORER TESTS
# ==========================================

def test_grant_scorer_and_eligibility():
    from grant_engine.grant_scorer import GrantScorer, ClientProfile
    
    scorer = GrantScorer()
    
    # 1. Test perfectly eligible and matching client
    client = ClientProfile(
        organization_name="PHMEG Solutions",
        organization_type="Nonprofit (501c3)",
        has_active_sam_registration=True,
        uei_number="UEI123456789",
        geographic_location="Washington, DC",
        cost_share_available=True,
        requested_budget=2500000.0,
        has_required_key_personnel=True
    )
    
    nofo = NOFOData(
        opportunity_number="HRSA-26-089",
        eligibility="Nonprofit organizations only.",
        award_ceiling="$4,000,000",
        award_floor="$100,000",
        cost_sharing="Yes, cost sharing is required for all applicants.",
        uei_sam_required="Yes"
    )
    
    res = scorer.score_eligibility(client, nofo)
    assert res["is_eligible"] is True
    assert res["score"] == 100
    assert len(res["disqualifications"]) == 0

    # 2. Test ineligible client (missing SAM, wrong entity type, over budget)
    ineligible_client = ClientProfile(
        organization_name="BadFit LLC",
        organization_type="Large For-Profit Corporation",
        has_active_sam_registration=False,
        uei_number="",
        cost_share_available=False,
        requested_budget=5000000.0,  # Exceeds ceiling
        has_required_key_personnel=False
    )
    
    res_bad = scorer.score_eligibility(ineligible_client, nofo)
    assert res_bad["is_eligible"] is False
    assert res_bad["score"] < 50
    assert len(res_bad["disqualifications"]) > 0
    # Confirm exact blockers are detected
    assert any("SAM.gov registration" in d for d in res_bad["disqualifications"])
    assert any("exceeds the award ceiling" in d for d in res_bad["disqualifications"])
    assert any("cost-sharing" in d for d in res_bad["disqualifications"])


def test_grants_gov_attachments_checklist():
    from grant_engine.validator import Validator
    from grant_engine.grant_context import GrantContext
    from grant_engine.template_loader import TemplateLoader
    
    with tempfile.TemporaryDirectory() as temp_dir:
        # Mini manifest and file mocking to prevent load exceptions
        manifest_data = {"templates": [{"id": "organizational_profile", "filename": "o.txt", "order": 1, "required_fields": []}]}
        with open(os.path.join(temp_dir, "manifest.json"), "w", encoding="utf-8") as f:
            json.dump(manifest_data, f)
        with open(os.path.join(temp_dir, "o.txt"), "w", encoding="utf-8") as f:
            f.write("Profile")
            
        loader = TemplateLoader(temp_dir)
        validator = Validator(loader)
        context = GrantContext()
        
        # Test empty context checklist (should report missing attachments)
        checklist = validator.check_grants_gov_attachments(context)
        assert len(checklist) == 5
        # Mandatory forms should be flagged as missing
        assert any(not item["is_present"] for item in checklist)
        assert any(item["status"] == "Missing Required Attachment" for item in checklist)


def test_semantic_alignment_matching():
    from grant_engine.grant_scorer import GrantScorer
    scorer = GrantScorer()
    
    client_caps = [
        "Our team provides enterprise cybersecurity, zero-trust firewalls, and active threat monitoring.",
        "We specialize in clinical healthcare delivery, primary care in rural clinics, and nursing staff support."
    ]
    
    requirements = [
        "Applicant must detail their strategy for rural primary care clinic staffing and healthcare access.",
        "Must support enterprise threat intelligence and zero-trust IT network security."
    ]
    
    alignments = scorer.semantic_align_capabilities(client_caps, requirements)
    assert len(alignments) == 2
    
    # First requirement should map to the healthcare capability statement (high score, aligned)
    assert "primary care" in alignments[0]["best_matching_client_capability"]
    assert alignments[0]["alignment_score"] > 30.0
    assert alignments[0]["is_aligned"] is True
    
    # Second requirement should map to the cybersecurity capability statement
    assert "cybersecurity" in alignments[1]["best_matching_client_capability"]
    assert alignments[1]["alignment_score"] > 30.0
    assert alignments[1]["is_aligned"] is True




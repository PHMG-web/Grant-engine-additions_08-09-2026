import re
from docx import Document
from typing import Dict, Any, List, Set, Union

class DocxExtractor:
    def __init__(self, docx_source: Any):
        """
        Initializes the DocxExtractor with a file path or a file-like object.
        Includes load error handling.
        """
        self.docx_path = docx_source if isinstance(docx_source, str) else "file-like object"
        try:
            if docx_source is None:
                raise ValueError("No DOCX source provided.")
            
            # Check if it is a FastAPI UploadFile
            if hasattr(docx_source, "file"):
                if hasattr(docx_source.file, "seek"):
                    docx_source.file.seek(0)
                self.doc = Document(docx_source.file)
            else:
                self.doc = Document(docx_source)
        except Exception as e:
            # Explicit failure reporting
            raise RuntimeError(f"Failed to load DOCX file: {str(e)}") from e

    def extract_full_text(self) -> str:
        """
        Extracts all paragraph and table text from the DOCX.
        """
        parts = []
        for para in self.doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)

        return "\n".join(parts)

    def _extract_text_from_runs(self, element: Any) -> str:
        """Helper to safely reconstruct text from consecutive runs to support split keys."""
        if hasattr(element, "runs") and element.runs:
            return "".join(run.text for run in element.runs)
        return getattr(element, "text", "")

    def _extract_matches_from_text(self, text: str, pattern: str, found: Set[str]):
        """Helper to scan text, match patterns, and clean brace placeholders."""
        matches = re.findall(pattern, text)
        for m in matches:
            clean_field = m.strip("{} ")
            if clean_field:
                found.add(clean_field)

    def extract_placeholders(self) -> List[str]:
        """
        Extracts all valid {field} (single or double curly braces) placeholders.
        Reconstructs run text to ensure split placeholders are detected.
        """
        found = set()
        pattern = r"\{+([a-zA-Z0-9_.-]+)\}+"

        # Paragraph placeholders
        for para in self.doc.paragraphs:
            text = self._extract_text_from_runs(para)
            self._extract_matches_from_text(text, pattern, found)

        # Table placeholders
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    if cell.paragraphs:
                        reconstructed = [self._extract_text_from_runs(p) for p in cell.paragraphs]
                        text = "\n".join(reconstructed) if any(reconstructed) else cell.text
                    self._extract_matches_from_text(text, pattern, found)

        return sorted(list(found))

    def detect_malformed_tokens(self) -> List[str]:
        """
        Scans all text blocks in the document (paragraphs and tables) to check for:
        - unmatched open brace '{'
        - unmatched close brace '}'
        - empty braces '{}' (including with spaces)
        Provides detailed issue reporting with context and location.
        """
        issues = []
        
        # Scan paragraphs
        for idx, para in enumerate(self.doc.paragraphs):
            text = self._extract_text_from_runs(para)
            para_issues = self._scan_text_for_malformed(text)
            for issue in para_issues:
                issues.append(f"Paragraph {idx+1}: {issue}")

        # Scan tables
        for t_idx, table in enumerate(self.doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    text = cell.text
                    if cell.paragraphs:
                        reconstructed = [self._extract_text_from_runs(p) for p in cell.paragraphs]
                        text = "\n".join(reconstructed) if any(reconstructed) else cell.text

                    cell_issues = self._scan_text_for_malformed(text)
                    for issue in cell_issues:
                        issues.append(f"Table {t_idx+1}, Row {r_idx+1}, Cell {c_idx+1}: {issue}")

        return issues

    def _scan_text_for_malformed(self, text: str) -> List[str]:
        """
        Helper method to scan a block of text for malformed braces.
        """
        issues = []
        
        # 1. Check for empty braces, e.g. {}, { }, {{}}
        empty_matches = re.finditer(r"\{+[ \t]*\}+", text)
        for m in empty_matches:
            issues.append(f"Empty placeholder '{m.group(0)}' found")

        # 2. Check for unmatched braces using stack
        stack = []
        for idx, char in enumerate(text):
            if char == "{":
                stack.append(idx)
            elif char == "}":
                if stack:
                    stack.pop()
                else:
                    # Unmatched closing brace
                    context = text[max(0, idx-15):min(len(text), idx+15)]
                    issues.append(f"Unmatched closing brace '}}' in context: '...{context.strip()}...'")
        
        # Any remaining on stack are unmatched opening braces
        for idx in stack:
            context = text[max(0, idx-15):min(len(text), idx+15)]
            issues.append(f"Unmatched opening brace '{{' in context: '...{context.strip()}...'")

        return issues

    def extract_all(self) -> Dict[str, Any]:
        """
        Extracts all text, placeholders, and detects any malformed issues.
        """
        return {
            "text": self.extract_full_text(),
            "placeholders": self.extract_placeholders(),
            "issues": self.detect_malformed_tokens()
        }

    def correct_text_braces(self, text: str) -> tuple[str, int]:
        """
        Automatically corrects simple unmatched and empty braces in a string.
        - '{field' (unmatched open) -> '{field}'
        - 'field}' (unmatched close) -> '{field}'
        - '{}' or '{   }' (empty) -> deleted
        Returns:
            (corrected_text, total_fixes_made)
        """
        if not text:
            return text, 0

        fixes = 0

        # 1. Clean/delete empty placeholders (e.g. {}, {   }, {{}}, etc.)
        cleaned_text, count = re.subn(r"\{+[ \t]*\}+", "", text)
        fixes += count

        # 2. Fix unmatched open braces:
        # Match pattern: '{FIELD' which is not followed by a closing '}' before another brace or end of sentence.
        unmatched_open_pattern = r"\{([a-zA-Z0-9_.-]+)(?![^{}]*\})"
        
        matches = list(re.finditer(unmatched_open_pattern, cleaned_text))
        # Process from back to front to avoid index shifting issues
        for m in reversed(matches):
            start, end = m.start(), m.end()
            field = m.group(1)
            cleaned_text = cleaned_text[:start] + "{" + field + "}" + cleaned_text[end:]
            fixes += 1

        # 3. Fix unmatched close braces:
        # Scan from left to right to find unmatched close braces '}'
        stack = []
        unmatched_close_indices = []
        for idx, char in enumerate(cleaned_text):
            if char == "{":
                stack.append(idx)
            elif char == "}":
                if stack:
                    stack.pop()
                else:
                    unmatched_close_indices.append(idx)

        # Process unmatched closing brace indices from back to front
        for idx in reversed(unmatched_close_indices):
            # Scan backwards to find the word boundary starting point
            word_start = idx
            while word_start > 0 and re.match(r"[a-zA-Z0-9_.-]", cleaned_text[word_start-1]):
                word_start -= 1
            
            if word_start < idx:
                word = cleaned_text[word_start:idx]
                cleaned_text = cleaned_text[:word_start] + "{" + word + "}" + cleaned_text[idx+1:]
                fixes += 1

        return cleaned_text, fixes

    def auto_correct_braces(self, output_path: str = None) -> int:
        """
        Scans and automatically corrects all unmatched and empty braces in paragraphs and table cells.
        Saves the corrected document to output_path if provided.
        Returns:
            The total number of fixes made across the entire document.
        """
        total_fixes = 0

        # 1. Correct standard paragraphs
        total_fixes += self._correct_doc_paragraphs()

        # 2. Correct table cells
        total_fixes += self._correct_doc_tables()

        if output_path and total_fixes > 0:
            self.doc.save(output_path)

        return total_fixes

    def _correct_doc_paragraphs(self) -> int:
        """Helper to iterate and correct doc paragraphs."""
        fixes_count = 0
        for para in self.doc.paragraphs:
            original_text = para.text
            corrected_text, fixes = self.correct_text_braces(original_text)
            if fixes > 0:
                para.text = corrected_text
                fixes_count += fixes
        return fixes_count

    def _correct_doc_tables(self) -> int:
        """Helper to iterate and correct doc table cell paragraphs."""
        fixes_count = 0
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    fixes_count += self._correct_single_cell(cell)
        return fixes_count

    def _correct_single_cell(self, cell: Any) -> int:
        """Helper to correct text of a single cell."""
        fixes_count = 0
        if cell.paragraphs:
            for para in cell.paragraphs:
                original_text = para.text
                corrected_text, fixes = self.correct_text_braces(original_text)
                if fixes > 0:
                    para.text = corrected_text
                    fixes_count += fixes
        else:
            original_text = cell.text
            corrected_text, fixes = self.correct_text_braces(original_text)
            if fixes > 0:
                cell.text = corrected_text
                fixes_count += fixes
        return fixes_count

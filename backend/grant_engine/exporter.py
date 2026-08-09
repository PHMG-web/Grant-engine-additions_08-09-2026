import json
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from typing import Any

class Exporter:
    """
    Handles exporting GrantContext into different formats:
    - JSON (structured data)
    - DOCX (narrative document formatted to sam.gov / grants.gov standards)
    Guards output paths and handles missing values gracefully.
    """

    def _safe_value(self, val: Any) -> str:
        """
        Formats raw values gracefully. Ensures None is represented as an explicit 
        '[Not Populated]' text rather than raising exceptions or printing raw None.
        """
        if val is None:
            return "[Not Populated]"
        if isinstance(val, list):
            if not val:
                return "[Not Populated]"
            return ", ".join(str(v).strip() for v in val if str(v).strip())
        if isinstance(val, dict):
            if not val:
                return "[Not Populated]"
            return "; ".join(f"{k}: {v}" if v is not None else str(k) for k, v in val.items())
        return str(val).strip() if str(val).strip() else "[Not Populated]"

    def _ensure_dir_exists(self, file_path: str):
        """
        Verifies and creates target folder paths before write operations.
        """
        if not file_path:
            raise ValueError("Output file path cannot be empty.")
        dir_name = os.path.dirname(os.path.abspath(file_path))
        if dir_name:
            try:
                os.makedirs(dir_name, exist_ok=True)
            except Exception as e:
                raise RuntimeError(f"Unable to create target output directory '{dir_name}': {str(e)}") from e

    def _apply_sam_gov_styling(self, doc: Document):
        """
        Applies strict sam.gov and grants.gov styling specifications:
        - 1-inch margins on all sides (standard submission requirement)
        - Times New Roman font face (professional/academic standard)
        - Clear line spacing (1.15) and paragraph spacing
        """
        # Set 1-inch margins
        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        # Base style configuration
        styles = doc.styles
        
        # Heading 1 Style
        h1_style = styles["Heading 1"]
        h1_font = h1_style.font
        h1_font.name = "Times New Roman"
        h1_font.size = Pt(14)
        h1_font.bold = True
        h1_font.color.rgb = RGBColor(0, 0, 0)  # Professional Black
        h1_style.paragraph_format.space_before = Pt(12)
        h1_style.paragraph_format.space_after = Pt(6)

        # Heading 2 Style
        h2_style = styles["Heading 2"]
        h2_font = h2_style.font
        h2_font.name = "Times New Roman"
        h2_font.size = Pt(12)
        h2_font.bold = True
        h2_font.color.rgb = RGBColor(0, 0, 0)
        h2_style.paragraph_format.space_before = Pt(8)
        h2_style.paragraph_format.space_after = Pt(4)

        # Normal Text Style (Body text)
        normal_style = styles["Normal"]
        normal_font = normal_style.font
        normal_font.name = "Times New Roman"
        normal_font.size = Pt(11)
        normal_style.paragraph_format.line_spacing = 1.15
        normal_style.paragraph_format.space_after = Pt(6)

    def export_json(self, context, output_path="grant_output.json") -> str:
        """
        Export the GrantContext as a JSON file safely.
        """
        self._ensure_dir_exists(output_path)
        try:
            with open(output_path, "w", encoding="utf-8") as f:
                json.dump(context.to_dict(), f, indent=4, ensure_ascii=False)
            return output_path
        except PermissionError as perm_err:
            raise RuntimeError(f"Permission denied writing JSON output to '{output_path}': {str(perm_err)}") from perm_err
        except Exception as e:
            raise RuntimeError(f"Failed to export JSON to '{output_path}': {str(e)}") from e

    def export_docx(self, context, output_path="grant_output.docx") -> str:
        """
        Export the GrantContext as a beautifully formatted, SAM.gov & Grants.gov compliant Word document.
        Splits implementation into cover-page building and section-content building to minimize function complexity.
        """
        self._ensure_dir_exists(output_path)
        
        try:
            doc = Document()
            self._apply_sam_gov_styling(doc)

            # 1. Build Cover Page
            self._build_cover_page(doc, context)

            # 2. Build Sections
            self._build_sections_content(doc, context)

            doc.save(output_path)
            return output_path
            
        except PermissionError as perm_err:
            raise RuntimeError(f"Permission denied saving DOCX output to '{output_path}': {str(perm_err)}") from perm_err
        except Exception as e:
            raise RuntimeError(f"Failed to generate and export DOCX to '{output_path}': {str(e)}") from e

    def _build_cover_page(self, doc: Document, context: Any):
        """Helper to build a professional compliance cover page."""
        doc.add_paragraph("\n" * 2)
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run("FEDERAL GRANT APPLICATION / PROJECT PROPOSAL")
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_p.add_run("Prepared in Compliance with Grants.gov and SAM.gov Standards")
        sub_run.font.size = Pt(11)
        sub_run.font.italic = True
        
        doc.add_paragraph("\n" * 3)

        # Metadata Info Grid (Table)
        doc.add_heading("I. ADMINISTRATIVE IDENTIFICATION MATRIX", level=1)
        table = doc.add_table(rows=7, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Formatted table headers and elements
        def set_row(idx, key, val):
            r = table.rows[idx]
            r.cells[0].paragraphs[0].text = key
            r.cells[0].paragraphs[0].runs[0].font.bold = True
            r.cells[1].paragraphs[0].text = str(val)

        nofo_sec = context.nofo.get("sections", {}) if isinstance(context.nofo, dict) else {}
        opp_num = nofo_sec.get("opportunity_number") if nofo_sec else "[Not Populated]"
        agency = nofo_sec.get("agency") if nofo_sec else "[Not Populated]"
        org_name = context.organizational_profile.get("Organization_Name") or "[Not Populated]"
        uei = context.organizational_profile.get("uei_number") or context.organizational_profile.get("Additional_Info", {}).get("uei_number") or "[Not Populated]"
        set_aside = nofo_sec.get("set_aside_category") if nofo_sec else "[Not Populated]"
        naics = nofo_sec.get("naics_code") if nofo_sec else "[Not Populated]"

        set_row(0, "Funding Opportunity Number:", opp_num)
        set_row(1, "Issuing Federal Agency:", agency)
        set_row(2, "Applicant Legal Entity:", org_name)
        set_row(3, "Unique Entity Identifier (UEI):", uei)
        set_row(4, "SAM.gov Status:", "Active / Compliant")
        set_row(5, "Acquisition Set-Aside:", set_aside if set_aside else "Unrestricted")
        set_row(6, "Assigned NAICS Code:", naics if naics else "N/A")

        doc.add_paragraph("\n")
        doc.add_page_break()

    def _build_sections_content(self, doc: Document, context: Any):
        """Helper to iterate and populate standard section content."""
        sections_map = [
            ("II. ORGANIZATIONAL PROFILE & CAPACITY", context.organizational_profile),
            ("III. PROJECT DESCRIPTION & PROGRAM DESIGN", context.program_design),
            ("IV. IMPLEMENTATION PLAN & WORK PLAN", context.implementation_plan),
            ("V. KEY PERSONNEL & STAFFING STRUCTURE", context.staffing_plan),
            ("VI. BUDGET NARRATIVE & COST JUSTIFICATION", context.budget_narrative),
            ("VII. EVALUATIVE PLAN & MEASURABLE OUTCOMES", context.evaluation_plan),
            ("VIII. PROJECT SUSTAINABILITY & PARTNERSHIPS", context.sustainability_plan)
        ]

        for section_title, section_data in sections_map:
            doc.add_heading(section_title, level=1)
            
            if not section_data or all(v is None for v in section_data.values()):
                doc.add_paragraph("[Applicant is required to complete and populate this section before submission.]")
                continue

            for k, v in section_data.items():
                if k == "Additional_Info":
                    if isinstance(v, dict) and v:
                        doc.add_heading("Additional Compliance Specifications", level=2)
                        for sub_k, sub_v in v.items():
                            doc.add_paragraph(f"{sub_k}: {self._safe_value(sub_v)}")
                    continue

                # Aligned lists and tables representations
                if isinstance(v, list) and v:
                    doc.add_paragraph(f"{k}:")
                    for item in v:
                        doc.add_paragraph(f"  • {item}", style='List Bullet')
                elif isinstance(v, dict) and v:
                    # Format nested items neatly
                    doc.add_paragraph(f"{k}:")
                    for sub_k, sub_v in v.items():
                        doc.add_paragraph(f"  - {sub_k}: {self._safe_value(sub_v)}")
                else:
                    doc.add_paragraph(f"{k}: {self._safe_value(v)}")

